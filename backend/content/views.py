from django.http import JsonResponse
import os
from io import BytesIO
from django.http import FileResponse
from rest_framework import status
from rest_framework.response import Response
from rest_framework.views import APIView
from rest_framework.permissions import IsAuthenticated
from accounts.services import log_activity, enforce_user_budget

from sites_app.models import AIPlatformSettings, Site

from .ai import generate_content, generate_product, get_provider_catalog
from .models import GeneratedContent
from assistant.models import AIUsageEvent
from .serializers import (
    GenerateProductSerializer,
    GenerateSerializer,
    GeneratedContentSerializer,
)


def site_from_key(request):
    key = request.headers.get("X-API-Key") or request.data.get("api_key")
    return Site.objects.filter(api_key=key).first() if key else None


def dashboard_site(request):
    sid = request.data.get("site_id") or request.query_params.get("site_id")
    # Dashboard requests must identify a Site explicitly. This prevents one
    # tenant's content from being generated/stored against another Site.
    return Site.objects.filter(id=sid).first() if sid else None


def record_ai_usage(request, site, usage, feature="content_generation"):
    if not usage:
        return
    AIUsageEvent.objects.create(
        user=request.user if getattr(request, "user", None) and request.user.is_authenticated else None,
        site=site,
        feature=feature,
        provider=usage.get("provider", ""),
        model=usage.get("model", ""),
        api_calls=int(usage.get("api_calls", 1) or 1),
        input_tokens=int(usage.get("input_tokens", 0) or 0),
        output_tokens=int(usage.get("output_tokens", 0) or 0),
        total_tokens=int(usage.get("total_tokens", 0) or 0),
        cost_usd=usage.get("cost_usd", 0),
        metadata=usage.get("metadata", {}),
    )


class GenerateView(APIView):
    permission_classes = [IsAuthenticated]
    """Universal content-generation endpoint used by the dashboard and embed."""

    def post(self, request):
        site = site_from_key(request) or dashboard_site(request)
        if not site:
            return Response(
                {"detail": "No Content Consult Site is configured."},
                status=status.HTTP_400_BAD_REQUEST,
            )

        enforce_user_budget(request.user)
        serializer = GenerateSerializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        data = serializer.validated_data
        # AI provider is selected only by an admin through the Site AI settings.
        data.pop("provider", None)

        try:
            result = generate_content(site=site, **data)
        except Exception as exc:
            return Response(
                {"detail": str(exc)},
                status=status.HTTP_502_BAD_GATEWAY,
            )

        usage = result.pop("_usage", None)
        item = GeneratedContent.objects.create(
            site=site,
            created_by=request.user,
            content_type=data["content_type"],
            topic=data["topic"],
            tone=data["tone"],
            word_count=data["word_count"],
            **result,
        )
        record_ai_usage(request, site, usage)
        log_activity(request.user, "generate", f"Generated {data['content_type']}: {data['topic']}", {"content_id": item.id, "site_id": site.id, "provider": usage.get("provider") if usage else None, "model": usage.get("model") if usage else None, "api_calls": usage.get("api_calls", 0) if usage else 0, "cost_usd": str(usage.get("cost_usd", 0)) if usage else "0"}, request)
        return Response(GeneratedContentSerializer(item).data, status=status.HTTP_201_CREATED)


class GenerateProductView(APIView):
    permission_classes = [IsAuthenticated]
    """Backward-compatible product endpoint."""

    def post(self, request):
        site = site_from_key(request) or dashboard_site(request)
        if not site:
            return Response(
                {"detail": "No Content Consult Site is configured."},
                status=status.HTTP_400_BAD_REQUEST,
            )

        enforce_user_budget(request.user)
        serializer = GenerateProductSerializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        data = serializer.validated_data
        data.pop("provider", None)

        try:
            result = generate_product(site=site, **data)
        except Exception as exc:
            return Response(
                {"detail": str(exc)},
                status=status.HTTP_502_BAD_GATEWAY,
            )

        usage = result.pop("_usage", None)
        item = GeneratedContent.objects.create(
            site=site,
            created_by=request.user,
            content_type="product",
            topic=data["topic"],
            tone=data["tone"],
            word_count=data["word_count"],
            **result,
        )
        record_ai_usage(request, site, usage)
        log_activity(request.user, "generate", f"Generated product: {data['topic']}", {"content_id": item.id, "site_id": site.id, "provider": usage.get("provider") if usage else None, "model": usage.get("model") if usage else None, "api_calls": usage.get("api_calls", 0) if usage else 0, "cost_usd": str(usage.get("cost_usd", 0)) if usage else "0"}, request)
        return Response(GeneratedContentSerializer(item).data, status=status.HTTP_201_CREATED)


class ContentHistoryView(APIView):
    permission_classes = [IsAuthenticated]
    def get(self, request):
        site = site_from_key(request) or dashboard_site(request)
        qs = GeneratedContent.objects.all()
        if site:
            qs = qs.filter(site=site)
        if not request.user.is_staff:
            qs = qs.filter(created_by=request.user)
        return Response(GeneratedContentSerializer(qs[:50], many=True).data)


class DraftView(APIView):
    permission_classes = [IsAuthenticated]
    throttle_classes = []

    def _queryset(self, request):
        site = site_from_key(request) or dashboard_site(request)
        qs = GeneratedContent.objects.filter(status="draft")
        if site:
            qs = qs.filter(site=site)
        if not request.user.is_staff:
            qs = qs.filter(created_by=request.user)
        return qs

    def get(self, request):
        return Response(GeneratedContentSerializer(self._queryset(request)[:100], many=True).data)

    def post(self, request):
        site = site_from_key(request) or dashboard_site(request)
        if not site:
            return Response({"detail": "No Content Consult Site is configured."}, status=status.HTTP_400_BAD_REQUEST)

        payload = request.data.copy()
        title = (payload.get("title") or payload.get("topic") or "Untitled draft").strip()
        body = payload.get("body") or ""
        metadata = payload.get("metadata") or {}
        if not isinstance(metadata, dict):
            metadata = {}
        for key in ("subject", "preheader", "cta", "platform"):
            if payload.get(key):
                metadata[key] = payload.get(key)
        # Saving generated content must be idempotent. AI generation already creates
        # the working content record so the user can review it; an explicit Save
        # updates that same record instead of creating a second identical draft.
        draft_id = payload.get("draft_id")
        item = None
        if draft_id:
            try:
                item = GeneratedContent.objects.get(
                    pk=draft_id,
                    site=site,
                    created_by=request.user,
                    status="draft",
                )
            except (GeneratedContent.DoesNotExist, ValueError, TypeError):
                return Response({"detail": "The generated draft no longer exists or is not owned by this user."}, status=status.HTTP_404_NOT_FOUND)

        if item is None:
            item = GeneratedContent.objects.create(
                site=site,
                created_by=request.user,
                content_type=payload.get("content_type", "blog"),
                topic=payload.get("topic", title),
                tone=payload.get("tone", "professional"),
                word_count=int(payload.get("word_count") or 300),
                title=title[:500],
                body=body,
                category=payload.get("category") or "",
                price=payload.get("price") if payload.get("price") not in ("", None) else None,
                seo_description=payload.get("seo_description") or "",
                keywords=payload.get("keywords") or [],
                metadata=metadata,
                status="draft",
            )
            created = True
        else:
            item.content_type = payload.get("content_type", item.content_type)
            item.topic = payload.get("topic", item.topic)
            item.tone = payload.get("tone", item.tone)
            item.word_count = int(payload.get("word_count") or item.word_count or 300)
            item.title = title[:500]
            item.body = body
            item.category = payload.get("category") or ""
            item.price = payload.get("price") if payload.get("price") not in ("", None) else None
            item.seo_description = payload.get("seo_description") or ""
            item.keywords = payload.get("keywords") or []
            item.metadata = metadata
            item.save()
            created = False

        log_activity(request.user, "save_draft", f"Saved draft: {title}", {"content_id": item.id, "site_id": site.id, "updated_existing": not created}, request)
        return Response(GeneratedContentSerializer(item).data, status=status.HTTP_201_CREATED if created else status.HTTP_200_OK)

    def delete(self, request):
        draft_id = request.query_params.get("id")
        if not draft_id:
            return Response({"detail": "Draft id is required."}, status=status.HTTP_400_BAD_REQUEST)
        qs = self._queryset(request)
        try:
            item = qs.get(pk=draft_id)
        except GeneratedContent.DoesNotExist:
            return Response({"detail": "Draft not found."}, status=status.HTTP_404_NOT_FOUND)
        title = item.title
        item.delete()
        log_activity(request.user, "delete_draft", f"Deleted draft: {title}", {"content_id": draft_id}, request)
        return Response(status=status.HTTP_204_NO_CONTENT)


class DraftExportView(APIView):
    permission_classes = [IsAuthenticated]

    def _get_draft(self, request, draft_id):
        qs = GeneratedContent.objects.filter(pk=draft_id, status="draft")
        if not request.user.is_staff:
            qs = qs.filter(created_by=request.user)
        return qs.first()

    def get(self, request, draft_id, fmt):
        item = self._get_draft(request, draft_id)
        if not item:
            return Response({"detail": "Draft not found."}, status=status.HTTP_404_NOT_FOUND)

        safe_title = "".join(c if c.isalnum() or c in " -_" else "_" for c in (item.title or "draft"))[:80].strip() or "draft"
        metadata = item.metadata if isinstance(item.metadata, dict) else {}

        if fmt == "word":
            from docx import Document
            document = Document()
            document.add_heading(item.title or "Untitled draft", 0)
            document.add_paragraph(f"Content type: {item.content_type.replace('_', ' ').title()}")
            for key, label in (("subject", "Subject"), ("preheader", "Preheader"), ("platform", "Platform"), ("cta", "CTA")):
                if metadata.get(key):
                    document.add_paragraph(f"{label}: {metadata[key]}")
            if item.body:
                for paragraph in str(item.body).split("\n"):
                    document.add_paragraph(paragraph)
            if item.seo_description:
                document.add_heading("SEO description", level=2)
                document.add_paragraph(item.seo_description)
            if item.keywords:
                document.add_heading("Keywords / hashtags", level=2)
                document.add_paragraph(" ".join(item.keywords))
            output = BytesIO()
            document.save(output)
            output.seek(0)
            return FileResponse(output, as_attachment=True, filename=f"{safe_title}.docx", content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

        if fmt == "pdf":
            from reportlab.lib.pagesizes import A4
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.enums import TA_LEFT
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
            from reportlab.lib.units import mm
            output = BytesIO()
            doc = SimpleDocTemplate(output, pagesize=A4, rightMargin=18*mm, leftMargin=18*mm, topMargin=18*mm, bottomMargin=18*mm, title=item.title or "Draft")
            styles = getSampleStyleSheet()
            body_style = ParagraphStyle("DraftBody", parent=styles["BodyText"], leading=17, spaceAfter=8, alignment=TA_LEFT)
            story = [Paragraph(item.title or "Untitled draft", styles["Title"]), Spacer(1, 8)]
            story.append(Paragraph(f"Content type: {item.content_type.replace('_', ' ').title()}", styles["Normal"]))
            for key, label in (("subject", "Subject"), ("preheader", "Preheader"), ("platform", "Platform"), ("cta", "CTA")):
                if metadata.get(key):
                    story.append(Paragraph(f"<b>{label}:</b> {str(metadata[key]).replace('&','&amp;').replace('<','&lt;').replace('>','&gt;')}", styles["Normal"]))
            story.append(Spacer(1, 10))
            import html
            for paragraph in str(item.body or "").split("\n"):
                text = html.escape(paragraph) if paragraph.strip() else "&nbsp;"
                story.append(Paragraph(text, body_style))
            if item.seo_description:
                story.append(Paragraph("SEO description", styles["Heading2"]))
                story.append(Paragraph(html.escape(item.seo_description), body_style))
            if item.keywords:
                story.append(Paragraph("Keywords / hashtags", styles["Heading2"]))
                story.append(Paragraph(html.escape(" ".join(item.keywords)), body_style))
            doc.build(story)
            output.seek(0)
            return FileResponse(output, as_attachment=True, filename=f"{safe_title}.pdf", content_type="application/pdf")

        return Response({"detail": "Unsupported export format."}, status=status.HTTP_400_BAD_REQUEST)


class SiteListView(APIView):
    permission_classes = [IsAuthenticated]
    def get(self, request):
        site = Site.objects.first()
        if not site:
            site = Site.objects.create(
                name=os.getenv("NIFTYBOT_SITE_NAME", "Content Consult"),
                domain=os.getenv("NIFTYBOT_SITE_DOMAIN", "localhost"),
                brand_voice=os.getenv("NIFTYBOT_BRAND_VOICE", "clear, professional and human"),
                language=os.getenv("NIFTYBOT_LANGUAGE", "English"),
            )
        return Response([
            {
                "id": site.id,
                "name": site.name,
                "domain": site.domain,
                "brand_voice": site.brand_voice,
                "language": site.language,
                "ai_provider": AIPlatformSettings.get_solo().active_provider,
                "ai_provider_label": {"gemini": "Gemini", "openai": "ChatGPT / OpenAI", "anthropic": "Claude", "xai": "Grok / xAI"}.get(AIPlatformSettings.get_solo().active_provider, "Gemini"),
                "ai_model": next((p["model"] for p in get_provider_catalog() if p["id"] == AIPlatformSettings.get_solo().active_provider), None),
            }
        ])


class ProviderCatalogView(APIView):
    """Return safe provider/model metadata without exposing API keys."""
    permission_classes = [IsAuthenticated]

    def get(self, request):
        return Response({
            "providers": [item for item in get_provider_catalog() if item["configured"]]
        })


class HealthView(APIView):
    """Lightweight readiness endpoint used by the Content Studio."""
    permission_classes = []

    def get(self, request):
        return JsonResponse({"ok": True, "service": "content-consult-django"})
